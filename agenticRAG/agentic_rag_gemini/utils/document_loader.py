"""Document loader for PDF, Word, and image files with OCR support.

Supports:
- PDF files (text extraction + OCR for scanned documents)
- Word files (.docx)
- Images (.png, .jpg, .jpeg, .gif, .bmp)
- Mixed content (text + images with OCR)
"""

import os
import io
from typing import List, Dict, Optional, Tuple
from pathlib import Path
import tempfile

from utils.logger import get_logger

logger = get_logger(__name__)


class DocumentLoader:
    """Load and extract text from multiple document formats."""
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg', '.gif', '.bmp'}
    
    def __init__(self):
        """Initialize document loader with available libraries."""
        self.has_pypdf = self._check_library('pypdf')
        self.has_docx = self._check_library('docx', 'python-docx')
        self.has_pytesseract = self._check_library('pytesseract')
        self.has_pdf2image = self._check_library('pdf2image')
        self.has_pil = self._check_library('PIL', 'pillow')
        
        logger.info(
            f"Document loader initialized - "
            f"PDF: {self.has_pypdf}, DOCX: {self.has_docx}, "
            f"OCR: {self.has_pytesseract}, PDF2Image: {self.has_pdf2image}, PIL: {self.has_pil}"
        )
    
    @staticmethod
    def _check_library(module_name: str, package_name: str = None) -> bool:
        """Check if a library is installed."""
        try:
            __import__(module_name)
            return True
        except ImportError:
            return False
    
    def load_file(self, file_path: str) -> str:
        """Load and extract text from a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format not supported
            FileNotFoundError: If file doesn't exist
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self._load_pdf(file_path)
        elif suffix in {'.docx', '.doc'}:
            return self._load_docx(file_path)
        elif suffix == '.txt':
            return self._load_text(file_path)
        elif suffix in {'.png', '.jpg', '.jpeg', '.gif', '.bmp'}:
            return self._load_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def load_directory(self, directory_path: str) -> Dict[str, str]:
        """Load all supported files from a directory.
        
        Args:
            directory_path: Path to directory
            
        Returns:
            Dictionary mapping file names to extracted text
        """
        results = {}
        dir_path = Path(directory_path)
        
        if not dir_path.is_dir():
            raise ValueError(f"Not a directory: {directory_path}")
        
        for file_path in dir_path.iterdir():
            if file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                try:
                    logger.info(f"Loading {file_path.name}...")
                    text = self.load_file(str(file_path))
                    results[file_path.name] = text
                    logger.info(f"Successfully loaded {file_path.name} ({len(text)} chars)")
                except Exception as e:
                    logger.error(f"Failed to load {file_path.name}: {e}")
        
        return results
    
    def _load_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file.
        
        Strategy:
        1. Extract text using pypdf (fast, handles text-based PDFs)
        2. Also render each page as image and OCR it (catches screenshots, figures, diagrams)
        3. Merge both: deduplicated text + OCR text from images
        """
        if not self.has_pypdf:
            raise ImportError("pypdf not installed. Install with: pip install pypdf")
        
        try:
            import pypdf
            
            text_parts = []
            
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                num_pages = len(reader.pages)
                
                logger.info(f"Extracting text from {file_path.name} ({num_pages} pages)")
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
            
            extracted_text = "\n".join(text_parts)
            
            # Always try OCR for image content (screenshots, figures, diagrams)
            # This catches text inside images that pypdf cannot extract
            ocr_text = self._ocr_pdf_images(file_path)
            if ocr_text:
                extracted_text = extracted_text + "\n\n" + ocr_text if extracted_text.strip() else ocr_text
            
            # Final fallback: if still no text, try full-page OCR
            if len(extracted_text.strip()) < 100:
                logger.info("Detected scanned PDF, attempting full-page OCR...")
                extracted_text = self._ocr_pdf(file_path)
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise
    
    def _ocr_pdf_images(self, file_path: Path) -> str:
        """Extract text from images embedded in PDF pages using OCR.
        
        Renders each page as an image and runs OCR to capture text in
        screenshots, figures, diagrams, and other visual elements.
        
        Returns:
            OCR-extracted text from PDF images, or empty string if unavailable
        """
        if not self.has_pytesseract or not self.has_pdf2image or not self.has_pil:
            return ""
        
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            logger.info(f"Running image OCR on PDF pages...")
            
            images = convert_from_path(str(file_path))
            ocr_parts = []
            
            for page_num, image in enumerate(images):
                try:
                    ocr_text = pytesseract.image_to_string(image)
                    if ocr_text and ocr_text.strip():
                        ocr_parts.append(f"--- Page {page_num + 1} (Image OCR) ---\n{ocr_text.strip()}")
                except Exception as e:
                    logger.warning(f"OCR failed on page {page_num + 1}: {e}")
            
            if ocr_parts:
                logger.info(f"OCR extracted text from {len(ocr_parts)} pages")
            
            return "\n".join(ocr_parts)
            
        except Exception as e:
            logger.warning(f"PDF image OCR failed: {e}")
            return ""
    
    def _ocr_pdf(self, file_path: Path) -> str:
        """Extract text from scanned PDF using OCR."""
        if not self.has_pytesseract or not self.has_pdf2image:
            logger.warning(
                "OCR requires pytesseract and pdf2image. "
                "Install with: pip install pytesseract pdf2image"
            )
            return "[PDF contains images only - OCR not available]"
        
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            logger.info(f"Running OCR on scanned PDF...")
            
            images = convert_from_path(str(file_path))
            text_parts = []
            
            for page_num, image in enumerate(images):
                try:
                    logger.info(f"OCR processing page {page_num + 1}...")
                    text = pytesseract.image_to_string(image)
                    if text.strip():
                        text_parts.append(f"--- Page {page_num + 1} (OCR) ---\n{text}")
                except Exception as e:
                    logger.warning(f"OCR failed on page {page_num + 1}: {e}")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"OCR processing failed: {e}")
            return "[PDF OCR processing failed]"
    
    def _load_docx(self, file_path: Path) -> str:
        """Extract text from Word document."""
        if not self.has_docx:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables if any
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_data))
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            # Extract text from embedded images if OCR is available
            if self.has_pytesseract and self.has_pil:
                try:
                    import pytesseract
                    from PIL import Image
                    
                    image_texts = []
                    for rel in doc.part._rels:
                        rel_obj = doc.part._rels[rel]
                        if "image" in rel_obj.target_ref:
                            try:
                                img = Image.open(io.BytesIO(rel_obj.target_part.blob))
                                ocr_text = pytesseract.image_to_string(img)
                                if ocr_text.strip():
                                    image_texts.append(ocr_text.strip())
                            except Exception as img_err:
                                logger.warning(f"Failed to OCR docx image: {img_err}")
                    
                    if image_texts:
                        text_parts.append("\n--- Embedded Image Content ---")
                        text_parts.extend(image_texts)
                        logger.info(f"Extracted OCR text from {len(image_texts)} embedded images in Word doc")
                except Exception as ocr_err:
                    logger.warning(f"DOCX OCR failed: {ocr_err}")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error reading Word document: {e}")
            raise
    
    def _load_text(self, file_path: Path) -> str:
        """Load plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def _load_image(self, file_path: Path) -> str:
        """Extract text from image file using OCR."""
        if not self.has_pytesseract:
            raise ImportError(
                "pytesseract not installed. Install with: pip install pytesseract"
            )
        
        if not self.has_pil:
            raise ImportError("pillow not installed. Install with: pip install pillow")
        
        try:
            from PIL import Image
            import pytesseract
            
            logger.info(f"Extracting text from image {file_path.name}...")
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            return text if text.strip() else "[Image contains no readable text]"
            
        except Exception as e:
            logger.error(f"Error reading image: {e}")
            raise
    
    def batch_load_files(
        self, 
        file_paths: List[str], 
        chunk_size: Optional[int] = None
    ) -> List[Dict[str, str]]:
        """Load multiple files and optionally chunk them.
        
        Args:
            file_paths: List of file paths
            chunk_size: If specified, split text into chunks of this size
            
        Returns:
            List of documents with metadata
        """
        documents = []
        
        for file_path in file_paths:
            try:
                text = self.load_file(file_path)
                
                if chunk_size:
                    chunks = self._chunk_text(text, chunk_size)
                    for chunk_num, chunk in enumerate(chunks):
                        documents.append({
                            'file': Path(file_path).name,
                            'chunk': chunk_num,
                            'text': chunk,
                            'type': 'document_chunk'
                        })
                else:
                    documents.append({
                        'file': Path(file_path).name,
                        'chunk': 0,
                        'text': text,
                        'type': 'document'
                    })
                    
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {e}")
        
        return documents
    
    @staticmethod
    def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start = end - overlap
        
        return chunks
    
    def load_mixed_content(self, directory: str) -> List[Dict[str, str]]:
        """Load all documents from directory with metadata.
        
        Args:
            directory: Path to directory with documents
            
        Returns:
            List of document dictionaries with content and metadata
        """
        documents = []
        dir_path = Path(directory)
        
        for file_path in sorted(dir_path.iterdir()):
            if file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                try:
                    text = self.load_file(str(file_path))
                    documents.append({
                        'file_name': file_path.name,
                        'file_path': str(file_path),
                        'content': text,
                        'file_type': file_path.suffix.lower(),
                        'size': len(text),
                        'timestamp': file_path.stat().st_mtime
                    })
                except Exception as e:
                    logger.error(f"Failed to load {file_path.name}: {e}")
        
        return documents
